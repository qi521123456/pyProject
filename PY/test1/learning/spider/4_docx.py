from docx import Document
from docx.shared import Inches
import requests
def readDocx(docName):
    fullText = []
    doc = Document(docName)
    paras = doc.paragraphs
    for p in paras:
        print(p)
        fullText.append(p.text)
    return '\n'.join(fullText)


def pa188():
    url = 'http://10.0.1.188:8081/sol_manager-0.0.1-SNAPSHOT/toolList'
    header = {
        'Cookie': 'JSESSIONID=53DC293CEC7DFD55CE585D7565D50D8E'
    }
    data = {'currentpage':'1'}
    r = requests.post(url,data=data,headers=header)
    print(r.text)

def writedocx(filename):
    recordset = [{'qty':1,'id':111,'desc':"love"}]
    document = Document()
    head_style = {}
    t = document.add_heading('一、关键信息基础设施联网状况和安全隐患', 1)
    t.style = "Heading2"
    t.add_run("楷体_gb2312").font.name = "仿宋"
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

   # document.add_picture('D:/location.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    hdr_cells[3].text = '大大'
    for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['qty'])
        row_cells[1].text = str(item['id'])
        row_cells[2].text = item['desc']
    #table.style = 'LightShading-Accent2'
    #document.add_page_break()

    document.save(filename)
def testdoc(filename):
    styles = ['Normal', 'Body Text', 'Body Text 2', 'Body Text 3', 'Caption', 'Heading 1',
              'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7',
              'Heading 8', 'Heading 9', 'Intense Quote', 'List', 'List 2', 'List 3', 'List Bullet',
              'List Bullet 2', 'List Bullet 3', 'List Continue', 'List Continue 2', 'List Continue 3',
              'List Number', 'List Number 2', 'List Number 3', 'List Paragraph',
              'No Spacing', 'Quote', 'Subtitle', 'TOCHeading', 'Title']
    styles2 = ['Body Text Char', 'Body Text 2 Char', 'Body Text 3 Char', 'Book Title',
               'Default Paragraph Font', 'Emphasis', 'Heading 1 Char', 'Heading 2 Char',
               'Heading 3 Char', 'Heading 4 Char', 'Heading 5 Char', 'Heading 6 Char',
               'Heading 7 Char', 'Heading 8 Char', 'Heading 9 Char', 'Intense Emphasis',
               'Intense Quote Char', 'Intense Reference', 'Macro Text Char', 'Quote Char',
               'Strong', 'Subtitle Char', 'Subtle Emphasis', 'Subtle Reference', 'Title Char']

    doc = Document()
    for s in styles:
        doc.add_paragraph('一、关键信息基础设施联网状况和安全隐患\t'+s, style=s)

    doc.save(filename)


def readtable(filename):
    doc = Document(filename)
    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
if __name__ == '__main__':
    #writedocx("D:/demo.docx")
    testdoc("D:/test.docx")

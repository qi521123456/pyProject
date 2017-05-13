# coding:utf-8
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook,load_workbook
from openpyxl.styles import Font,colors
import matplotlib.pyplot as plt
from pylab import mpl
import pymysql
import datetime,os
ICDTABLES = [{'table_name': 'result_iec104', 'vendor': '未知', 'device_info':None},
             {'table_name': 'result_melsec_q_tcp', 'vendor': 'Mitsubishi', 'device_info':'cpuinfo'},
             {'table_name': 'result_melsec_q_udp', 'vendor': 'Mitsubishi', 'device_info':'cpuinfo'},
             {'table_name': 'result_modbus', 'vendor_name': 'vendor', 'vendor': None, 'device_info':'device'},
             {'table_name': 'result_moxa', 'vendor_name': 'vendor', 'vendor': None, 'device_info':'server_name'},
             {'table_name': 'result_s7comm', 'vendor_name': 'vendor', 'vendor': None, 'device_info':'basic_hardware'},
             {'table_name': 'result_snmp', 'vendor_name': 'vendor', 'vendor': None, 'device_info':'cpu'},
             {'table_name': 'result_fox', 'vendor_name': 'brand_id', 'vendor': None, 'device_info':None},
             {'table_name': 'result_fins', 'vendor': 'Omron', 'device_info':'controller_model'},
             {'table_name': 'result_ethernetip', 'vendor_name': 'vendor', 'vendor': None, 'device_info':'product_name'},
             {'table_name': 'result_bacnet', 'vendor_name': 'vendor_name', 'vendor': None, 'device_info':'object_name'},
             {'table_name': 'result_cspv4', 'vendor': '未知', 'device_info':None},
             {'table_name': 'result_dnp3', 'vendor': '未知', 'device_info':None}
             ]
IofTABLES = ['result_http', 'result_dahua']
PROVINCES = ['上海市', '云南', '内蒙古自治区', '北京市', '台湾', '吉林', '四川', '天津市', '宁夏回族自治区', '安徽',
             '山东', '山西', '广东', '广西壮族自治区', '新疆', '新疆维吾尔自治区', '江苏', '江西', '河北', '河南',
             '浙江', '海南', '湖北', '湖南', '澳门特别行政区', '甘肃', '福建', '西藏自治区', '贵州', '辽宁', '重庆市',
             '陕西', '青海', '香港特别行政区', '黑龙江']

def get_data(province,begin_time,end_time):  # province为汉字省份，time为2017-3-4-10
    con = pymysql.connect(host='10.0.1.199',
                          port=3306,
                          user='root',
                          password='123456',
                          db='sol_daily',
                          charset='utf8mb4',
                          cursorclass=pymysql.cursors.DictCursor)
    bt = [int(b) for b in begin_time.split('-')]
    et = [int(e) for e in end_time.split('-')]
    begin = datetime.datetime(bt[0],bt[1],bt[2])
    end = datetime.datetime(et[0],et[1],et[2])
    icdTables = ICDTABLES
    iofTables = IofTABLES
    icdData = []
    iofData = []
    try:
        with con.cursor() as cursor:
            ipsicd = []
            for table in icdTables:
                table_name = table['table_name']
                vendor = table['vendor']
                info = table['device_info']
                if vendor is None and info is None:
                    q_sql = "SELECT device_ip,device_city,device_port,create_time,%s as vendor FROM %s " \
                            "WHERE device_province='%s'" % (table['vendor_name'],table_name,province)
                elif vendor is not None and info is None:
                    q_sql = "SELECT device_ip,device_city,device_port,create_time FROM %s " \
                            "WHERE device_province='%s'" % (table_name,province)
                elif vendor is None and info is not None:
                    q_sql = "SELECT device_ip,device_city,device_port,create_time,%s as vendor,%s as info FROM %s " \
                            "WHERE device_province='%s'" % (table['vendor_name'],info,table_name,province)
                else:
                    q_sql = "SELECT device_ip,device_city,device_port,create_time,%s as info FROM %s " \
                            "WHERE device_province='%s'" % (info,table_name,province)
                cursor.execute(q_sql)
                tmp = cursor.fetchall()
                if tmp is None:
                    continue
                for data in tmp:
                    d = {}
                    # print(str(year) + '-' + str(month),data['create_time'].strftime('%Y-%m'))
                    timestamp = data['create_time']
                    if timestamp < begin or timestamp > end:
                        continue
                    d['ip'] = data['device_ip']
                    d['location'] = province + data['device_city']
                    d['port'] = data['device_port']
                    if vendor is None:
                        vendor = data['vendor']
                    d['vendor'] = vendor
                    d['info'] = "None"
                    if info is not None:
                        d['info'] = data['info']
                    p_sql = "SELECT protocol_name FROM protocol WHERE protocol_port='%s'" % d['port']
                    cursor.execute(p_sql)
                    protocol = cursor.fetchone()
                    if protocol is None:
                        proto = "未知"
                    else:
                        proto = protocol['protocol_name']
                    d['protocol'] = proto
                    if d['ip'] in ipsicd:
                        continue
                    ipsicd.append(d['ip'])
                    icdData.append(d)
            ipsiof = []
            for table in iofTables:
                q_sql = "SELECT create_time,device_ip,device_city,device_port,device FROM %s " \
                        "WHERE device_province='%s'" % (table,province)
                cursor.execute(q_sql)
                tmp = cursor.fetchall()
                if tmp is None:
                    continue
                for data in tmp:
                    d = {}
                    #print(str(year) + '-' + str(month),data['create_time'].strftime('%Y-%m'))
                    timestamp = data['create_time']
                    if timestamp < begin or timestamp > end:
                        continue
                    d['ip'] = data['device_ip']
                    d['location'] = province+data['device_city']
                    d['port'] = data['device_port']
                    d['vendor'] = data['device']
                    p_sql = "SELECT protocol_name FROM protocol WHERE protocol_port='%s'" % d['port']
                    cursor.execute(p_sql)
                    protocol = cursor.fetchone()
                    if protocol is None:
                        proto = "未知"
                    else:
                        proto = protocol['protocol_name']
                    d['protocol'] = proto
                    if d['ip'] in ipsiof:
                        continue
                    ipsiof.append(d['ip'])
                    iofData.append(d)
    finally:
        con.close()
    return icdData,iofData
def draw_pie(info,province,path):
    location = {}
    vendor = {}
    protocol = {}
    if len(info) <= 0:
        print("no data")
        return
    flag = 'info' in info[0]
    tmp = {'location':location,'vendor':vendor,'protocol':protocol}
    for i in info:
        for k,v in tmp.items():
            if i[k] in v:
                v[i[k]] += 1
            else:
                v[i[k]] = 1
    for name,data in tmp.items():
        mpl.rcParams['font.sans-serif'] = ['SimHei']  # 指定默认字体 解决中文显示问题
        plt.figure()
        if name == 'location':
            if flag:
                title = province+"物联网设备地区分布图"
            else:
                title = province+"联网核心工控设备地区分布图"
        elif name == 'vendor':
            if flag:
                title = province + "物联网设备厂商分布图"
            else:
                title = province+"联网核心工控设备厂商分布图"
        else:
            if flag:
                title = province + "物联网设备协议分布图"
            else:
                title = province+"联网核心工控设备协议分布图"
        plt.title(title)
        labels = []
        sizes = []
        for k,v in data.items():
            labels.append(k)
            sizes.append(v)
        #colors = ['red', 'yellowgreen', 'lightskyblue']
        plt.pie(sizes, labels=labels,autopct='%3.1f%%', shadow=False,startangle=90, pctdistance=0.5)
        plt.axis('equal')
        plt.legend(loc='upper right')
        # plt.show()
        plt.savefig(path+name+".png")
        plt.close()
def draw_grid(data,province,path):
    if len(data) <= 0:
        print("no data")
        return
    name1 = "所有设备"
    name2 = "有信息设备"

    wb = Workbook()
    ws = wb.active
    #ws = wb.create_sheet(name1)
    ws.title = name1
    ft = Font(name='Courier New', size=12, color=colors.BLACK, bold=True)
    ws['A1'].value = 'IP'
    ws['B1'].value = '协议'
    ws['C1'].value = '端口'
    ws['D1'].value = '位置'
    ws['E1'].value = '厂商'
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 25
    #ws.column_dimensions['F'].width = 25
    ws.column_dimensions['A'].width = 25
    ws['A1'].font = ft
    ws['B1'].font = ft
    ws['C1'].font = ft
    ws['D1'].font = ft
    ws['E1'].font = ft
    for row,d in enumerate(data):
        #for col in range(0,5):
        ws.cell(column=1,row=row+2,value=d['ip'])
        ws.cell(column=2, row=row + 2, value=d['protocol'])
        ws.cell(column=3, row=row + 2, value=d['port'])
        ws.cell(column=4, row=row + 2, value=d['location'])
        ws.cell(column=5, row=row + 2, value=d['vendor'])
    if 'info' in data[0]:
        infos = []
        for i in data:
            if (i['info'] == 'None') or (i['info'].strip() == '') or (i['info'] == 'xxxxxx'):
                continue
            infos.append(i)
        ws2 = wb.create_sheet(name2)
        ws2.column_dimensions['B'].width = 25
        ws2.column_dimensions['C'].width = 25
        ws2.column_dimensions['D'].width = 25
        ws2.column_dimensions['E'].width = 25
        ws2.column_dimensions['A'].width = 25
        ws2['A1'].value = 'IP'
        ws2['B1'].value = '协议'
        ws2['C1'].value = '位置'
        ws2['D1'].value = '厂商'
        ws2['E1'].value = '设备信息'
        ws2['A1'].font = ft
        ws2['B1'].font = ft
        ws2['C1'].font = ft
        ws2['D1'].font = ft
        ws2['E1'].font = ft
        for row, d in enumerate(infos):
            ws2.cell(column=1, row=row + 2, value=d['ip'])
            ws2.cell(column=2, row=row + 2, value=d['protocol'])
            ws2.cell(column=3, row=row + 2, value=d['location'])
            ws2.cell(column=4, row=row + 2, value=d['vendor'])
            ws2.cell(column=5, row=row + 2, value=d['info'])
    wb.save(path+province+'.xlsx')

def create_doc(data,province,filename,begin_time,end_time):
    b = begin_time.split('-')
    e = end_time.split('-')
    path = filename[:filename.rfind('/')+1]
    count = len(data)
    if count <= 0:
        print("no data")
        return
    doc = Document()
    doc.add_heading("一、关键信息基础设施联网状况和安全隐患", 1)
    doc.add_heading("1.联网关键信息基础设施探测分析", 2)
    if 'info' in data[0]:
        doc.add_paragraph("针对%s的联网关键信息基础设施开展主动探测与识别分析，范围覆盖常用的工业协议及端口，例如S7comm、Modbus、BACnet、Fox等。" % province)
        p = doc.add_paragraph("自%s年%s月%s日开始，截止至%s年%s月%s日，" % (b[0],b[1],b[2],e[0],e[1],e[2]))
        p.add_run("一共探测发现联网关键信息基础设施核心工业设备%s例" % count).bold = True
        p.add_run("。安徽联网关键信息基础设施资产明细详见附件“%s.xslx”。%s联网关键信息基础设施资产统计分析图如下所示。" % (province,province))
        doc.add_picture(path+'location.png',width=Inches(4))
        doc.add_picture(path + 'protocol.png',width=Inches(4))
        doc.add_picture(path + 'vendor.png',width=Inches(4))
        doc.add_heading("2.安全漏洞分析", 2)
        doc.add_paragraph("探测到的部分联网核心设备和监控系统存在漏洞缺陷，具有潜在的安全风险。一旦遭受网络攻击，将造成重大危害和损失，建议对联网的关键信息基础设施及早处置。")
        doc.add_heading("2.1.联网核心智能设备的漏洞匹配", 3)
        doc.add_paragraph("目前探测到的安徽地区的核心智能设备主要分为两类，一类是目标主机回复了S7comm等的响应，但不包含具体的设备信息，"
                          "只能判断为关键信息基础设施核心智能设备；另一类则包含有详细的设备信息，其中设备详单信息见附件“%s.xlsx”。" % province)
        doc.add_paragraph("通过将设备型号与漏洞库做匹配，能够获知相应联网核心智能设备可能存在的安全漏洞，经过统计，安徽地区存在漏洞的"
                          "设备集中分布在罗克韦尔设备，对应的设备型号存在的漏洞信息如下表所示。（具体设备IP详见%s.xlsx）。" % province)
        table = doc.add_table(rows=1, cols=1, style='Table Grid')
        table.rows[0].cells[0].text = province+"部分联网核心智能设备可能存在的安全漏洞..."
    else:
        doc.add_paragraph("针对%s的联网关键信息基础设施开展主动探测与识别分析，范围覆盖常用的物联网协议及端口，例如HTTP协议。" % province)
        p = doc.add_paragraph("自%s年%s月%s日开始，截止至%s年%s月%s日，" % (b[0], b[1], b[2], e[0], e[1], e[2]))
        p.add_run("一共探测发现联网关键信息基础设施核心工业设备%s例" % count).bold = True
        p.add_run("主要包括雄迈、海康威视和大华产品供应商。具体资产详情见附件“%s.xslx”。%s联网监控系统统计分布信息如下图所示。" % (province,province))
        doc.add_picture(path + 'location.png', width=Inches(4))
        doc.add_picture(path + 'vendor.png', width=Inches(4))
    doc.save(filename)

def main(path,province = "北京市",begin_time = '2017-03-22'):
    end_time = datetime.datetime.now().strftime('%Y-%m-%d')
    path += province+'/'
    if province not in PROVINCES:
        print("no province ‘%s’" % province)
        return
    if os.path.exists(path):
        __import__('shutil').rmtree(path)
    os.makedirs(path)
    icd,iof = get_data(province,begin_time,end_time)
    doc_type = {"工控":icd, "物联网":iof}
    for d_t,data in doc_type.items():
        os.mkdir(path + d_t + "/")
        draw_pie(data,province,path+d_t + "/")
        draw_grid(data, province, path + d_t + "/")
        create_doc(data,province,path+d_t + "/"+province+"联网关键信息基础建设-"+d_t+'.docx',begin_time,end_time)


if __name__ == '__main__':
    main("D:/")


